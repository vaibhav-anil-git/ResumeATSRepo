from docx import Document
from io import BytesIO

def txt_to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
